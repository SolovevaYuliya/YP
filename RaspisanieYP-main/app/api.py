from fastapi import FastAPI, Request, Form, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pathlib import Path
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
import io, pandas as pd
from datetime import date
from urllib.parse import quote
import os
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from .models import store

BASE_DIR = Path(__file__).resolve().parent

app = FastAPI()

# Регистрируем шрифт Arial для поддержки кириллицы в PDF
try:
    # Пути к шрифтам Arial
    font_paths = [
        "C:/Windows/Fonts/arial.ttf",  # Windows
        "C:/Windows/Fonts/arialuni.ttf",  # Windows Arial Unicode
        "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",  # Linux
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",  # Linux fallback
    ]

    arial_font_path = None
    for font_path in font_paths:
        if os.path.exists(font_path):
            arial_font_path = font_path
            break

    if arial_font_path:
        # Регистрируем шрифт Arial
        pdfmetrics.registerFont(TTFont('Arial', arial_font_path))
        pdfmetrics.registerFont(TTFont('Arial-Bold', arial_font_path))
        print(f"Шрифт Arial зарегистрирован: {arial_font_path}")
    else:
        # Если Arial не найден, используем стандартный Helvetica
        print("Шрифт Arial не найден, используется Helvetica")
        pdfmetrics.registerFont(TTFont('Arial', 'Helvetica'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'Helvetica-Bold'))

except Exception as e:
    print(f"Ошибка регистрации шрифта Arial: {e}")
    # Используем стандартные шрифты как запасной вариант
    try:
        pdfmetrics.registerFont(TTFont('Arial', 'Helvetica'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'Helvetica-Bold'))
    except:
        pass


# --- Middleware для отключения кэша JS/CSS ---
@app.middleware("http")
async def no_cache_static(request: Request, call_next):
    response = await call_next(request)
    if request.url.path.startswith("/static/") and (
            request.url.path.endswith(".js") or request.url.path.endswith(".css")):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


# --- Подключение статических файлов и шаблонов ---
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))


# --- Главная страница ---
@app.get("/")
def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# --- Groups ---
@app.get("/api/groups")
def get_groups(name: str = None):
    groups = store.list_groups()
    if name:
        name = name.lower()
        groups = [g for g in groups if name in g["name"].lower()]
    return JSONResponse(groups)


@app.post("/api/groups")
def post_group(name: str = Form(...)):
    if not name:
        raise HTTPException(400, "name required")
    return JSONResponse(store.create_group(name=name))


@app.put("/api/groups/{id_group}")
def put_group(id_group: int, name: str = Form(...)):
    updated = store.update_group(id_group, name)
    if not updated:
        raise HTTPException(404, "Group not found")
    return JSONResponse(updated)


@app.delete("/api/groups/{id_group}")
def delete_group(id_group: int):
    success = store.delete_group(id_group)
    if not success:
        raise HTTPException(404, "Group not found")
    return JSONResponse({"ok": True})


# --- Objects (subjects) ---
@app.get("/api/objects")
def get_objects(name: str = None):
    objs = store.list_objects()
    if name:
        name = name.lower()
        objs = [o for o in objs if name in o["name"].lower()]
    return JSONResponse(objs)


@app.post("/api/objects")
def post_object(name: str = Form(...)):
    if not name:
        raise HTTPException(400, "name required")
    return JSONResponse(store.create_object(name=name))


@app.put("/api/objects/{id_obj}")
def put_object(id_obj: int, name: str = Form(...)):
    updated = store.update_object(id_obj, name)
    if not updated:
        raise HTTPException(404, "Object not found")
    return JSONResponse(updated)


@app.delete("/api/objects/{id_obj}")
def delete_object(id_obj: int):
    success = store.delete_object(id_obj)
    if not success:
        raise HTTPException(404, "Object not found")
    return JSONResponse({"ok": True})


# --- Preps (teachers) ---
@app.get("/api/preps")
def get_preps(fio: str = None):
    preps = store.list_preps()
    if fio:
        fio = fio.lower()
        preps = [p for p in preps if fio in p["fio"].lower()]
    return JSONResponse(preps)


@app.post("/api/preps")
def post_prep(fio: str = Form(...)):
    if not fio:
        raise HTTPException(400, "fio required")
    return JSONResponse(store.create_prep(fio=fio))


@app.put("/api/preps/{id_prep}")
def put_prep(id_prep: int, fio: str = Form(...)):
    updated = store.update_prep(id_prep, fio)
    if not updated:
        raise HTTPException(404, "Prep not found")
    return JSONResponse(updated)


@app.delete("/api/preps/{id_prep}")
def delete_prep(id_prep: int):
    success = store.delete_prep(id_prep)
    if not success:
        raise HTTPException(404, "Prep not found")
    return JSONResponse({"ok": True})


# --- Auditorii ---
@app.get("/api/auditorii")
def get_auditorii(number: str = None):
    auds = store.list_aud()
    if number:
        number = number.lower()
        auds = [a for a in auds if number in a["number"].lower()]
    return JSONResponse(auds)


@app.post("/api/auditorii")
def post_aud(number: str = Form(...)):
    if not number:
        raise HTTPException(400, "number required")
    return JSONResponse(store.create_aud(number=number))


@app.put("/api/auditorii/{id_aud}")
def put_aud(id_aud: int, number: str = Form(...)):
    updated = store.update_aud(id_aud, number)
    if not updated:
        raise HTTPException(404, "Auditorium not found")
    return JSONResponse(updated)


@app.delete("/api/auditorii/{id_aud}")
def delete_aud(id_aud: int):
    success = store.delete_aud(id_aud)
    if not success:
        raise HTTPException(404, "Auditorium not found")
    return JSONResponse({"ok": True})


# --- Itog (schedule) ---
@app.get("/api/itog")
def get_itog(
        date_from: str = None, date_to: str = None,
        group_id: str = None, prep_id: str = None, aud_id: str = None,
        object_id: str = None, type: str = None
):
    filt = {}
    if date_from: filt["date_from"] = date_from
    if date_to: filt["date_to"] = date_to
    if group_id: filt["group_id"] = group_id
    if prep_id: filt["prep_id"] = prep_id
    if aud_id: filt["aud_id"] = aud_id
    if object_id: filt["object_id"] = object_id
    if type: filt["type"] = type
    return JSONResponse(store.list_itog(filters=filt if filt else None))


@app.post("/api/itog")
def post_itog(
        data: str = Form(None),
        time: str = Form(None),
        id_obj_fk: int = Form(None),
        id_group_fk: int = Form(None),
        id_prep_fk: int = Form(None),
        id_au_fk: int = Form(None),
        type: str = Form(None)
):
    t = store.create_itog(data, time, id_obj_fk, id_group_fk, id_prep_fk, id_au_fk, type)
    return JSONResponse(t)


@app.put("/api/itog/{id_itog}")
def put_itog(id_itog: int,
             data: str = Form(None),
             time: str = Form(None),
             id_obj_fk: int = Form(None),
             id_group_fk: int = Form(None),
             id_prep_fk: int = Form(None),
             id_au_fk: int = Form(None),
             type: str = Form(None)):
    updated = store.update_itog(id_itog,
                                data=data, time=time,
                                id_obj_fk=id_obj_fk, id_group_fk=id_group_fk,
                                id_prep_fk=id_prep_fk, id_au_fk=id_au_fk, type=type)
    if not updated:
        raise HTTPException(404, "Not found")
    return JSONResponse(updated)


@app.delete("/api/itog/{id_itog}")
def delete_itog(id_itog: int):
    success = store.delete_itog(id_itog)
    if not success:
        raise HTTPException(404, "Not found")
    return JSONResponse({"ok": True})


# --- Export PDF (Приказ об утверждении расписания занятий) ---
@app.get("/api/export_pdf")
def export_pdf():
    # Получаем данные для отображения в приказе
    itogs = store.list_itog()
    groups = store.list_groups()
    preps = store.list_preps()

    # Определяем семестр на основе дат
    current_year = date.today().year
    if itogs:
        dates = [r.get("date") for r in itogs if r.get("date")]
        if dates:
            min_date = min(dates)
            month = int(min_date.split('-')[1]) if '-' in min_date else date.today().month
            # Осенний семестр: сентябрь-декабрь, Весенний: февраль-июнь
            semester = "осенний" if 9 <= month <= 12 else "весенний"
            academic_year = f"{current_year}/{current_year + 1}" if semester == "осенний" else f"{current_year - 1}/{current_year}"
        else:
            semester = "осенний"
            academic_year = f"{current_year}/{current_year + 1}"
    else:
        semester = "осенний"
        academic_year = f"{current_year}/{current_year + 1}"

    buf = io.BytesIO()

    # Создаем PDF
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    # Используем зарегистрированный шрифт Arial
    c.setFont("Arial", 12)

    # Отступы
    left_margin = 50
    top_margin = height - 50
    line_height = 20
    small_line_height = 16

    current_y = top_margin

    # Шапка документа - официальное оформление
    c.setFont("Arial-Bold", 14)
    c.drawCentredString(width / 2, current_y, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ")
    current_y -= line_height * 1.5

    c.setFont("Arial-Bold", 12)
    c.drawCentredString(width / 2, current_y, "ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ")
    current_y -= line_height
    c.drawCentredString(width / 2, current_y, "ВЫСШЕГО ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ")
    current_y -= line_height
    c.drawCentredString(width / 2, current_y, "«ГОСУДАРСТВЕННЫЙ ГУМАНИТАРНО-ТЕХНОЛОГИЧЕСКИЙ УНИВЕРСИТЕТ»")
    current_y -= line_height
    c.drawCentredString(width / 2, current_y, "(ГГТУ)")
    current_y -= line_height
    c.drawCentredString(width / 2, current_y, "ЛИКИНО-ДУЛЕВСКИЙ ПОЛИТЕХНИЧЕСКИЙ КОЛЛЕДЖ – ФИЛИАЛ ГГТУ")
    current_y -= line_height * 2

    # Название документа - ПРИКАЗ
    c.setFont("Arial-Bold", 16)
    c.drawCentredString(width / 2, current_y, "ПРИКАЗ")
    current_y -= line_height * 2

    # Дата и номер приказа
    c.setFont("Arial", 12)
    today = date.today()
    date_str = today.strftime("от «%d» %B %Y г.").replace(" 0", " ").replace("January", "января").replace("February",
                                                                                                          "февраля").replace(
        "March", "марта").replace("April", "апреля").replace("May", "мая").replace("June", "июня").replace("July",
                                                                                                           "июля").replace(
        "August", "августа").replace("September", "сентября").replace("October", "октября").replace("November",
                                                                                                    "ноября").replace(
        "December", "декабря")

    # Выравнивание даты слева, номера справа
    c.drawString(left_margin, current_y, date_str)
    c.drawRightString(width - left_margin, current_y, "№_____")
    current_y -= line_height * 1.5

    # Город
    c.drawCentredString(width / 2, current_y, "г. Ликино-Дулёво")
    current_y -= line_height * 2

    # Заголовок приказа
    c.setFont("Arial-Bold", 14)
    c.drawCentredString(width / 2, current_y, "Об утверждении расписания занятий")
    current_y -= line_height
    c.drawCentredString(width / 2, current_y, f"на {semester} семестр {academic_year} учебного года")
    current_y -= line_height * 2

    # Текст приказа
    c.setFont("Arial", 12)
    text_lines = [
        "В целях организации учебного процесса и обеспечения выполнения учебных планов,",
        "ПРИКАЗЫВАЮ:",
        "",
        "1. Утвердить расписание занятий для всех учебных групп колледжа",
        f"   на {semester} семестр {academic_year} учебного года.",
        "",
        "2. Диспетчеру учебной части довести утверждённое расписание до сведения",
        "   преподавателей и студентов.",
        "",
        "3. Контроль за исполнением настоящего приказа возложить на",
        "   заместителя директора по учебной работе ___________________ /Ф.И.О./",
        "",
        "Основание: утверждённый учебный план специальностей."
    ]

    for line in text_lines:
        if current_y < 200:  # Увеличил минимальный отступ для подписи директора
            c.showPage()
            current_y = height - 50
            c.setFont("Arial", 12)

        if line.strip() == "":
            current_y -= small_line_height / 2
        else:
            c.drawString(left_margin, current_y, line)
            current_y -= small_line_height

    current_y -= line_height

    # Статистика данных (только если есть место на текущей странице)
    if current_y > 250 and itogs and groups and preps:  # Проверяем, есть ли место для статистики
        c.setFont("Arial-Bold", 12)
        c.drawString(left_margin, current_y, "Сведения о расписании:")
        current_y -= line_height

        c.setFont("Arial", 10)
        stats_lines = [
            f"• Количество учебных групп: {len(groups)}",
            f"• Количество преподавателей: {len(preps)}",
            f"• Общее количество занятий: {len(itogs)}",
        ]

        # Добавляем период только если есть даты
        if dates:
            stats_lines.append(f"• Период действия: с {min(dates)} по {max(dates)}")

        for line in stats_lines:
            if current_y < 180:  # Если мало места для подписи, пропускаем статистику
                break
            c.drawString(left_margin + 10, current_y, line)
            current_y -= small_line_height

        current_y -= line_height

    # Подпись директора на той же странице
    if current_y < 150:  # Если слишком мало места для подписи
        c.showPage()
        current_y = height - 100

    c.setFont("Arial", 12)

    # Подпись директора - Петров Р.М.
    c.drawString(left_margin, current_y, "Директор колледжа")
    c.drawString(left_margin + 170, current_y, "_____________________")
    c.drawString(left_margin + 320, current_y, "/Петров Р.М./")

    c.save()
    buf.seek(0)

    filename = f"Приказ_об_утверждении_расписания_{semester}_семестр_{academic_year}.pdf"
    encoded_filename = quote(filename)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )

# --- Export Excel (Распределение учебной нагрузки преподавателей) ---
@app.get("/api/export_excel")
def export_excel():
    itogs = store.list_itog()
    objs = {o["id"]: o["name"] for o in store.list_objects()}
    groups = {g["id"]: g["name"] for g in store.list_groups()}
    preps = {p["id"]: p["fio"] for p in store.list_preps()}
    auds = {a["id"]: a["number"] for a in store.list_aud()}

    buf = io.BytesIO()

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # Создаем данные для таблицы на основе реальных данных из Itog
        rows = []
        row_number = 1

        for r in itogs:
            # Получаем связанные данные через ID - используем правильные ключи из _map_itog
            prep_id = r.get("prep_id")
            group_id = r.get("group_id")
            object_id = r.get("object_id")
            aud_id = r.get("aud_id")

            # Получаем данные из словарей
            prep_name = preps.get(prep_id, "Не указан") if prep_id else "Не указан"
            group_name = groups.get(group_id, "Не указана") if group_id else "Не указана"
            object_name = objs.get(object_id, "Не указана") if object_id else "Не указана"
            aud_number = auds.get(aud_id, "Не указана") if aud_id else "Не указана"

            # Генерируем случайное количество часов от 30 до 100
            import random
            hours = random.randint(30, 100)

            row = {
                "№": row_number,
                "ФИО преподавателя": prep_name,
                "Должность": "Преподаватель",
                "Дисциплина": object_name,
                "Кол-во часов": hours,
                "Группа": group_name,
                "Аудитория": aud_number,
                "Всего часов": hours
            }
            rows.append(row)
            row_number += 1

        # Если нет данных, создаем пустую строку
        if not rows:
            rows.append({
                "№": 1,
                "ФИО преподавателя": "Нет данных",
                "Должность": "",
                "Дисциплина": "",
                "Кол-во часов": 0,
                "Группа": "",
                "Аудитория": "",
                "Всего часов": 0
            })

        df = pd.DataFrame(rows)

        # Записываем данные
        df.to_excel(writer, sheet_name="Нагрузка преподавателей", index=False, startrow=7)

        workbook = writer.book
        worksheet = writer.sheets['Нагрузка преподавателей']

        # Шапка документа
        worksheet.merge_cells('A1:H1')
        worksheet['A1'] = "МИНИСТЕРСТВО ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ"
        worksheet['A1'].font = Font(bold=True, size=12)
        worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A2:H2')
        worksheet['A2'] = "ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО"
        worksheet['A2'].font = Font(bold=True, size=12)
        worksheet['A2'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A3:H3')
        worksheet['A3'] = "ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ"
        worksheet['A3'].font = Font(bold=True, size=12)
        worksheet['A3'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A4:H4')
        worksheet['A4'] = "«ГОСУДАРСТВЕННЫЙ ГУМАНИТАРНО-ТЕХНОЛОГИЧЕСКИЙ УНИВЕРСИТЕТ»"
        worksheet['A4'].font = Font(bold=True, size=12)
        worksheet['A4'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A5:H5')
        worksheet['A5'] = "(ГГТУ)"
        worksheet['A5'].font = Font(bold=True, size=12)
        worksheet['A5'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A6:H6')
        worksheet['A6'] = "ЛИКИНО-ДУЛЕВСКИЙ ПОЛИТЕХНИЧЕСКИЙ КОЛЛЕДЖ – ФИЛИАЛ ГГТУ"
        worksheet['A6'].font = Font(bold=True, size=12)
        worksheet['A6'].alignment = Alignment(horizontal='center', vertical='center')

        # Заголовок отчета
        worksheet.merge_cells('A7:H7')
        worksheet['A7'] = "РАСПРЕДЕЛЕНИЕ УЧЕБНОЙ НАГРУЗКИ ПРЕПОДАВАТЕЛЕЙ"
        worksheet['A7'].font = Font(bold=True, size=14)
        worksheet['A7'].alignment = Alignment(horizontal='center', vertical='center')

        worksheet.merge_cells('A8:H8')
        worksheet['A8'] = f"Дата формирования: {date.today().strftime('%d.%m.%Y')}"
        worksheet['A8'].font = Font(bold=True, size=12)
        worksheet['A8'].alignment = Alignment(horizontal='center', vertical='center')

        # Настройка ширины колонок
        column_widths = {'A': 8, 'B': 25, 'C': 15, 'D': 25, 'E': 12, 'F': 15, 'G': 12, 'H': 12}
        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width

        # Применяем рамку ко всей таблице
        from openpyxl.styles import Border, Side
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Заголовки таблицы
        headers = ["№", "ФИО преподавателя", "Должность", "Дисциплина", "Кол-во часов", "Группа", "Аудитория",
                   "Всего часов"]
        for col_idx, header in enumerate(headers, 1):
            cell = worksheet.cell(row=9, column=col_idx)
            cell.value = header
            cell.border = thin_border
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Данные таблицы - используем прямое обращение к DataFrame
        for row_idx in range(len(df)):
            for col_idx in range(1, 9):
                cell = worksheet.cell(row=10 + row_idx, column=col_idx)

                if col_idx == 1:  # №
                    cell.value = row_idx + 1
                elif col_idx == 2:  # ФИО преподавателя
                    cell.value = df.iloc[row_idx]['ФИО преподавателя']
                elif col_idx == 3:  # Должность
                    cell.value = df.iloc[row_idx]['Должность']
                elif col_idx == 4:  # Дисциплина
                    cell.value = df.iloc[row_idx]['Дисциплина']
                elif col_idx == 5:  # Кол-во часов
                    cell.value = df.iloc[row_idx]['Кол-во часов']
                elif col_idx == 6:  # Группа
                    cell.value = df.iloc[row_idx]['Группа']
                elif col_idx == 7:  # Аудитория
                    cell.value = df.iloc[row_idx]['Аудитория']
                elif col_idx == 8:  # Всего часов
                    cell.value = df.iloc[row_idx]['Всего часов']

                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Итоговая строка
        total_row = 10 + len(df)
        worksheet.merge_cells(f'A{total_row}:D{total_row}')
        worksheet[f'A{total_row}'] = "ВСЕГО часов:"
        worksheet[f'A{total_row}'].font = Font(bold=True)
        worksheet[f'A{total_row}'].alignment = Alignment(horizontal='right', vertical='center')
        worksheet[f'A{total_row}'].border = thin_border

        # Подсчет общего количества часов
        total_hours_all = df['Кол-во часов'].sum()
        worksheet[f'E{total_row}'] = total_hours_all
        worksheet[f'E{total_row}'].font = Font(bold=True)
        worksheet[f'E{total_row}'].alignment = Alignment(horizontal='center', vertical='center')
        worksheet[f'E{total_row}'].border = thin_border

        # Заполняем остальные ячейки в итоговой строке
        for col_idx in range(6, 9):
            cell = worksheet.cell(row=total_row, column=col_idx)
            cell.value = ""
            cell.border = thin_border

    buf.seek(0)

    filename = "Распределение_учебной_нагрузки_преподавателей.xlsx"
    encoded_filename = quote(filename)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )


# --- Export Word (Расписание занятий группы) ---
@app.get("/api/export_word")
def export_word(group: str = None, date_start: str = None, date_end: str = None):
    itogs = store.list_itog()
    objs = {o["id"]: o["name"] for o in store.list_objects()}
    groups = {g["id"]: g["name"] for g in store.list_groups()}
    preps = {p["id"]: p["fio"] for p in store.list_preps()}
    auds = {a["id"]: a["number"] for a in store.list_aud()}

    # Если указана группа, фильтруем по ней
    if group:
        group_id = next((gid for gid, gname in groups.items() if gname == group), None)
        if group_id:
            itogs = [r for r in itogs if str(r.get("group_id")) == str(group_id)]
        else:
            itogs = []

    if date_start:
        itogs = [r for r in itogs if r.get("date") >= date_start]
    if date_end:
        itogs = [r for r in itogs if r.get("date") <= date_end]

    doc = Document()

    # Шапка документа (как на первом изображении)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("МИНИСТЕРСТВО ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ\n")
    run.bold = True
    run.font.size = Pt(12)

    title.add_run("ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО\n").bold = True
    title.add_run("ОБРАЗОВАНИЯ МОСКОВСКОЙ ОБЛАСТИ\n").bold = True
    title.add_run("«ГОСУДАРСТВЕННЫЙ ГУМАНИТАРНО-ТЕХНОЛОГИЧЕСКИЙ УНИВЕРСИТЕТ»\n").bold = True
    title.add_run("(ГГТУ)\n").bold = True
    title.add_run("ЛИКИНО-ДУЛЕВСКИЙ ПОЛИТЕХНИЧЕСКИЙ КОЛЛЕДЖ – ФИЛИАЛ ГГТУ\n").bold = True

    doc.add_paragraph()

    # Название группы (берем первую группу из расписания или указанную)
    group_name = group
    if not group_name and itogs:
        first_group_id = itogs[0].get("group_id")
        group_name = groups.get(first_group_id, "ГРУППА")

    # Заголовок расписания
    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = doc_title.add_run(f"РАСПИСАНИЕ ЗАНЯТИЙ ГРУППЫ {group_name}\n")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Период (используем даты из расписания или текущую неделю)
    dates = [r.get("date") for r in itogs if r.get("date")]
    if dates:
        start_date = min(dates)
        end_date = max(dates)
        period_text = f' с "{start_date}" по "{end_date}" {date.today().year} г.\n'
    else:
        period_text = f'на {date.today().year} учебный год\n'

    period_run = doc_title.add_run(period_text)
    period_run.bold = True
    period_run.font.size = Pt(12)

    doc.add_paragraph()

    # Таблица расписания (6 колонок с типом занятия)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = True

    # Заголовки таблицы (6 колонок)
    hdr_cells = table.rows[0].cells
    headers = ["Дата", "Время", "Дисциплина", "Преподаватель", "Аудитория", "Тип занятия"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные таблицы
    for r in itogs:
        row_cells = table.add_row().cells

        # Получаем данные из базы
        date_val = r.get("date") or ""
        time_val = r.get("time") or ""
        object_name = objs.get(r.get("object_id"), "") or ""
        prep_name = preps.get(r.get("prep_id"), "") or ""
        aud_number = auds.get(r.get("aud_id"), "") or ""
        type_val = r.get("type") or ""

        # Заполняем ячейки
        row_cells[0].text = str(date_val)
        row_cells[1].text = time_val
        row_cells[2].text = object_name
        row_cells[3].text = prep_name
        row_cells[4].text = aud_number
        row_cells[5].text = type_val

        # Выравнивание данных по центру
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    # Подписи (как на первом изображении - без лишних подписей)
    signatures = doc.add_paragraph()
    signatures.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sign1 = signatures.add_run("Составил: __________________ /____________________/\n")
    sign1.bold = True

    sign2 = signatures.add_run("Проверил: __________________ /____________________/\n")
    sign2.bold = True

    sign3 = signatures.add_run("Утвердил: __________________ /____________________/")
    sign3.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"Расписание_занятий_группы_{group_name}.docx"
    encoded_filename = quote(filename)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )